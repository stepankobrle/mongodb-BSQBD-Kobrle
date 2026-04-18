"""
Generuje dokumentace.docx primo bez pouziti sablony z MD souboru.
Spusteni: python dokumentace/generuj_dokumentaci.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT = os.path.join(os.path.dirname(__file__), "dokumentace.docx")

# ── Barvy ──────────────────────────────────────────────────────────────
C_H1   = RGBColor(0x1F, 0x49, 0x7D)
C_H2   = RGBColor(0x2E, 0x74, 0xB5)
C_H3   = RGBColor(0x37, 0x86, 0x6F)
C_CODE = RGBColor(0xC7, 0x25, 0x4E)
C_BG   = "F4F4F4"
C_TH   = "D6E4F0"

doc = Document()

# ── Nastaveni okraju ────────────────────────────────────────────────────
for sec in doc.sections:
    sec.top_margin    = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin   = Cm(2.5)
    sec.right_margin  = Cm(2.0)

# ── Zakladni font ───────────────────────────────────────────────────────
for s in ("Normal", "Heading 1", "Heading 2", "Heading 3", "Heading 4"):
    if s in doc.styles:
        doc.styles[s].font.name = "Calibri"

# ── Pomocne funkce ──────────────────────────────────────────────────────
def shade_cell(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)

def shade_para(para, fill):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)

def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.size = Pt(16)
    p.runs[0].font.color.rgb = C_H1
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.size = Pt(13)
    p.runs[0].font.color.rgb = C_H2
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)

def h3(text):
    p = doc.add_heading(text, level=3)
    p.runs[0].font.size = Pt(11)
    p.runs[0].font.color.rgb = C_H3
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(1)

def para(text, space_after=4):
    p = doc.add_paragraph()
    _inline(p, text)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    _inline(p, text)
    p.paragraph_format.left_indent = Cm(0.4 + level * 0.5)
    p.paragraph_format.space_after = Pt(1)

def _inline(p, text):
    """Zpracuje **tucny** a `kod` v textu."""
    import re
    for part in re.split(r'(\*\*[^*]+\*\*|`[^`]+`)', text):
        if part.startswith("**") and part.endswith("**"):
            r = p.add_run(part[2:-2]); r.bold = True
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(part[1:-1])
            r.font.name = "Courier New"; r.font.size = Pt(8.5)
            r.font.color.rgb = C_CODE
        else:
            p.add_run(part)

def code(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    shade_para(p, C_BG)
    r = p.add_run(text)
    r.font.name = "Courier New"
    r.font.size  = Pt(7.5)

def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.autofit = True
    hcells = t.rows[0].cells
    for i, h in enumerate(headers):
        hcells[i].text = h
        hcells[i].paragraphs[0].runs[0].bold = True
        hcells[i].paragraphs[0].runs[0].font.size = Pt(9)
        shade_cell(hcells[i], C_TH)
    for row in rows:
        rc = t.add_row().cells
        for i, val in enumerate(row):
            rc[i].text = str(val)
            rc[i].paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def sep():
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(1)

def query(n, title, task, code_text, comment):
    h3(f"Dotaz {n}: {title}")
    para(f"**Úloha:** {task}", space_after=2)
    code(code_text)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    _inline(p, f"**Komentář:** {comment}")

# ══════════════════════════════════════════════════════════════════════════
# TITULNÍ STRANA
# ══════════════════════════════════════════════════════════════════════════
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_before = Pt(60)
r = tp.add_run("Semestrální práce – BSQBD")
r.bold = True; r.font.size = Pt(22); r.font.color.rgb = C_H1

tp2 = doc.add_paragraph()
tp2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = tp2.add_run("NoSQL dokumentová databáze – MongoDB 8.0 Sharded Cluster")
r2.font.size = Pt(14); r2.italic = True

doc.add_paragraph()
tp3 = doc.add_paragraph()
tp3.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp3.add_run("Student: Kobrle Štěpán  |  FEI  |  BSQBD – Big Data\nDatum: 2026")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════
# ÚVOD
# ══════════════════════════════════════════════════════════════════════════
h1("Úvod")
para("Semestrální práce se zabývá návrhem, nasazením a použitím **NoSQL dokumentové databáze MongoDB** v režimu **sharded clusteru s replikací**. Cílem je praktická ukázka horizontálního škálování (sharding), vysoké dostupnosti (replikace), agregační analytiky a zabezpečení.")
para("**Verze MongoDB: 8.0** (splňuje podmínku max. 3 verze zpět od aktuální).")
table(
    ["Docker obraz", "Zdroj", "Účel"],
    [
        ["`mongo:8.0`", "Docker Hub (oficiální)", "Datové nody, config servery, mongos routery"],
        ["`haohanyang/compass-web`", "Docker Hub (community)", "Webové UI MongoDB Compass (port 3000)"],
        ["`alpine:3.19`", "Docker Hub (oficiální)", "Generátor keyfile přes openssl"],
        ["`python:3.11-slim`", "Docker Hub (oficiální)", "Import dat (PyMongo + Pandas)"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════
# 1. ARCHITEKTURA
# ══════════════════════════════════════════════════════════════════════════
h1("1. Architektura")

h2("1.1 Schéma architektury")
para("Cluster se skládá z **13 Docker kontejnerů** ve čtyřech vrstvách:")
bullet("**Config Replica Set (configReplSet):** 3 nody (configsvr1–3) – uchovávají metadata shardingu.")
bullet("**Shard Replica Sets:** 3 shardy × 3 nody = 9 datových nodů (každý 1 PRIMARY + 2 SECONDARY).")
bullet("**Mongos Router Layer:** 2 routery (mongos1 port 27117, mongos2 port 27118) – query routing.")
bullet("**Klientská vrstva:** MongoCompass Web (port 3000), mongosh CLI, Python/PyMongo.")
para("Schéma je k dispozici v souboru `dokumentace/obrazky/schema_architektury.svg`.", space_after=6)

h2("1.2 Specifika konfigurace")

h3("1.2.1 CAP teorém")
para("MongoDB poskytuje garanci **CP (Consistency + Partition Tolerance)**.")
bullet("**Consistency:** Zápisy s `writeConcern: majority` jsou potvrzeny po replikaci na min. 2 nody. Čtení z PRIMARY garantuje linearizability.")
bullet("**Partition Tolerance:** Při síťovém rozdělení shard funguje, pokud PRIMARY má dostupnou majoritu (2/3 nodů).")
bullet("**Availability:** Omezená – během Raft election (10–30 s) je replica set dočasně nedostupný pro zápisy.")
para("**Zdůvodnění:** Filmová databáze je čtecí workload, kde konzistentní data (ratingy, metadata) jsou důležitější než krátkodobý výpadek při election. AP alternativa (Cassandra) by byla nevhodná.", space_after=6)

h3("1.2.2 Cluster")
para("**Počet clusterů: 1.** Projekt obsahuje ~110 000 dokumentů, objem v řádu stovek MB. Více clusterů by bylo relevantní pro multi-region deployment nebo workload isolation (OLTP vs. OLAP), což není cílem tohoto projektu.", space_after=6)

h3("1.2.3 Uzly")
table(
    ["Role", "Počet", "Konkrétní nody"],
    [
        ["Config server (metadata)", "3", "configsvr1, configsvr2, configsvr3"],
        ["Shard data node", "9", "shard{1,2,3}svr{1,2,3}"],
        ["Mongos router", "2", "mongos1, mongos2"],
    ]
)
para("3 config servery zajišťují HA metadata (majority = 2/3). 3 nody na každý shard jsou minimum pro automatickou election. 3 shardy zaručí smysluplnou distribuci dat (~33 % na shard).", space_after=6)

h3("1.2.4 Sharding / Partitioning")
table(
    ["Kolekce", "Shard key", "Typ", "Důvod"],
    [
        ["`movies`", "`{ id: \"hashed\" }`", "Hashed", "Unikátní TMDB id, vysoká kardinalita"],
        ["`credits`", "`{ movie_id: \"hashed\" }`", "Hashed", "Vazba na movies.id, rovnoměrné rozložení"],
        ["`ratings`", "`{ userId: \"hashed\" }`", "Hashed", "671 unikátních uživatelů × 100k+ záznamů"],
    ]
)
para("**Proč hashed:** Data nemají přirozený range pattern. Hashed eliminuje hot-spoty a zaručuje rovnoměrné rozložení. Nevýhoda: range dotazy jsou scatter-gather přes všechny shardy.")

para("**Sekundární indexy** (vytvořeny v `06-init-indexes.sh`):")
bullet("**movies:** `{ title: 1 }`, `{ release_date: 1 }`, `{ vote_average: -1 }`, `{ revenue: -1 }`, `{ budget: -1 }`, `{ title: \"text\", overview: \"text\" }` (fulltext)")
bullet("**credits:** `{ movie_id: 1 }` – klíčový pro JOIN operace přes `$lookup`")
bullet("**ratings:** `{ movieId: 1 }`, `{ userId: 1 }`, `{ rating: -1 }`")

h3("1.2.5 Replikace")
para("**Replikační faktor: 3** (1 PRIMARY + 2 SECONDARY) na každém shardu i config replica setu.")
bullet("**Majority quorum (fault tolerance = 1):** 2 ze 3 nodů – cluster toleruje výpadek 1 nodu bez ztráty dat.")
bullet("**Election bezpečnost:** Pro zvolení nového PRIMARY jsou potřeba 2 hlasy (minority se nikdy nestane PRIMARY).")
bullet("**Write concern:** `majority` – zápis potvrzen po replikaci na min. 2 nody. **Read preference:** `primary`.")
para("", space_after=4)

h3("1.2.6 Perzistence dat")
bullet("**Journaling (WAL):** Každý zápis jde nejprve do journal logu, poté do datových struktur. Synchronizace max. každých 100 ms. Garantuje recovery po pádu.")
bullet("**Checkpointing:** WiredTiger vytváří snapshot každých 60 sekund. Recovery = poslední checkpoint + journal.")
bullet("**Replikace jako záloha:** Při `writeConcern: majority` jsou data na min. 2 nodech – pád PRIMARY nezpůsobí ztrátu.")
bullet("**Docker volumes:** Každý node má dedikovaný named volume. Data přežijí `docker compose restart`, smazána při `docker compose down -v`.")
bullet("**WiredTiger cache:** 50 % RAM, komprese snappy (~3× redukce místa).")

h3("1.2.7 Distribuce dat")
para("**Zápisový tok:** klient → mongos → hash(shard key) → PRIMARY shardu → oplog → SECONDARY → potvrzení klientovi po majority.")
para("**Čtecí tok:** targeted query (shard key v `$match`) → 1 shard; scatter-gather (bez shard key) → všechny 3 shardy paralelně.")
table(
    ["Kolekce", "Shard 1", "Shard 2", "Shard 3", "Celkem"],
    [
        ["movies", "~1 600 dok.", "~1 600 dok.", "~1 600 dok.", "4 803"],
        ["credits", "~1 600 dok.", "~1 600 dok.", "~1 600 dok.", "4 803"],
        ["ratings", "~33 300 dok.", "~33 300 dok.", "~33 300 dok.", "100 004"],
    ]
)

h3("1.2.8 Zabezpečení")
bullet("**SCRAM-SHA-256:** Hesla uložena jako solené SHA-256 hashe, autentizace přes challenge-response.")
bullet("**RBAC:** `admin` (role `root`) pro správu clusteru, `filmuser` (role `readWrite` na filmdb) pro aplikaci.")
bullet("**Keyfile (interní auth mezi nody):** Generován dynamicky přes `openssl rand -base64 756` v kontejneru `keyfile-generator` (alpine). Nikdy není statický ani v gitu – každé spuštění má unikátní keyfile.")

# ══════════════════════════════════════════════════════════════════════════
# 2. FUNKČNÍ ŘEŠENÍ
# ══════════════════════════════════════════════════════════════════════════
h1("2. Funkční řešení")

h2("2.1 docker-compose.yml")
table(
    ["Skupina", "Služby", "Porty (host)"],
    [
        ["keyfile-generator", "alpine – vygeneruje keyfile", "–"],
        ["Config servery", "configsvr1–3", "–"],
        ["Shard 1", "shard1svr1–3", "–"],
        ["Shard 2", "shard2svr1–3", "–"],
        ["Shard 3", "shard3svr1–3", "–"],
        ["Mongos routery", "mongos1, mongos2", "27117, 27118"],
        ["mongo-init", "Init sekvence (01–06 skripty)", "–"],
        ["data-import", "Python import CSV dat", "–"],
        ["compass-web", "MongoCompass Web UI", "3000"],
    ]
)
para("Sdílený volume `init_data` slouží jako synchronizační mechanismus: flagy `/init-data/keyfile`, `/init-data/configreplset-ready` a `/init-data/phase2` řídí dvoufázovou inicializaci (Phase 1: bez --keyFile, Phase 2: s --keyFile).")
para("`mongo-init` kontejner spustí sekvenci: `01-init-configsvr.sh → 02-init-shards.sh → 03-init-mongos.sh → 04-init-users.sh → 05-init-validation.sh → 06-init-indexes.sh`. Celá inicializace probíhá automaticky v rámci jediného `docker compose up -d`.")

h2("2.2 Spuštění")
code("# Spuštění celého clusteru (automatická inicializace ~2–3 minuty)\ndocker compose up -d\n\n# Sledování průběhu inicializace\ndocker logs -f mongo-init\n\n# Přihlášení do clusteru (admin)\ndocker exec -it mongos1 mongosh -u admin -p adminpass123 --authenticationDatabase admin\n\n# Zastavení (zachovat data) / smazání dat\ndocker compose down\ndocker compose down -v")
para("MongoCompass Web: `http://localhost:3000`. Connection string: `mongodb://admin:adminpass123@mongos1:27017/admin`")

# ══════════════════════════════════════════════════════════════════════════
# 3. PŘÍPADY UŽITÍ
# ══════════════════════════════════════════════════════════════════════════
h1("3. Případy užití a případové studie")

h2("3.1 Obecné případy užití MongoDB")
for item in [
    "**Content Management Systems (CMS)** – heterogenní data s proměnlivým schématem.",
    "**Real-time analytika** – agregační pipeline přímo v databázi, bez ETL do data warehouse.",
    "**E-commerce katalogy** – produkty s různými atributy v jedné kolekci bez JOIN tabulek.",
    "**IoT a event logging** – vysoký throughput zápisů, time-series kolekce.",
    "**Personalizace a doporučování** – user profil s nested preferencemi v jednom dokumentu.",
]:
    bullet(item)

h2("3.2 Proč MongoDB a ne jiná NoSQL databáze")
bullet("**Redis** (klíč-hodnota) – výborný pro cache, ale neumí nested dokumenty ani multi-stage agregace.")
bullet("**Cassandra** (sloupcová) – AP systém, optimalizovaný pro write-heavy workloady. Chybí runtime JOIN a ad-hoc dotazování.")
bullet("**Elasticsearch** – výborný pro fulltext, ale chybí ACID transakce a striktní schéma validace.")

h2("3.3 Případová studie 1: eBay – Katalog produktů")
para("**Kontext:** eBay přešel na MongoDB od roku 2012 pro správu katalogu produktů a personalizovaných doporučení. MySQL katalog nedokázal efektivně pojmout heterogenní atributy napříč kategoriemi (elektronika vs. oblečení).")
para("**Řešení:** Každý produkt jako dokument s vnořenými atributy specifickými pro kategorii. Sharding podle `userId` a `productId` se stovkami shardů.")
para("**Výsledky:** Latence p99 klesla z 450 ms na 45 ms (10×). Počet tabulek klesl z 200+ na 12 kolekcí. Problém: counter documents → hot shard. Řešení: Redis in-memory counter + batch flush.", space_after=6)

h2("3.4 Případová studie 2: The Weather Channel – IoT ingestion")
para("**Kontext:** TWC obsluhuje 5 miliard meteorologických dotazů denně. Oracle RAC degradoval při špičkách (hurikány). Migrace na MongoDB Atlas 2014.")
para("**Řešení:** Time-series dokumenty ze stanic (každých 15 s), shard key `{ stationId: \"hashed\", timestamp: 1 }`. Replikační faktor 5 přes 3 datacentra.")
para("**Výsledky:** Throughput z 40 tisíc na 200 tisíc ingestí/sec. Náklady o 75 % nižší vs Oracle. Poučení: monotonní `timestamp` jako prefix shard key vytváří hot shard – `stationId` musí být první.", space_after=6)

h2("3.5 Případová studie 3: MetLife – Customer 360")
para('**Kontext:** MetLife měla zákaznická data v 70+ systémech. Operátor centra potřeboval 4+ minuty pro odpověď na "jaké máte produkty?"')
para("**Řešení:** MongoDB jako Customer 360 data hub. Každý klient = 1 dokument se všemi policies, claims a interakcemi. ETL z Oracle/DB2/Informix přes CDC. Shard key `{ customerId: \"hashed\" }`.")
para("**Výsledky:** Response time z 4 min na <5 s. 60 milionů klientů, 12 TB dat. Projekt dodán za 90 dní.")

# ══════════════════════════════════════════════════════════════════════════
# 4. VÝHODY A NEVÝHODY
# ══════════════════════════════════════════════════════════════════════════
h1("4. Výhody a nevýhody")

h2("Obecné výhody MongoDB")
for v in [
    "**Flexibilní schéma** – dokumenty v jedné kolekci mohou mít různé tvary.",
    "**Nested dokumenty** – 1:N vztahy bez JOIN tabulek (pole, objekty nativně v BSON).",
    "**Horizontální škálování (sharding)** – integrovaný, transparentní, automatický balancer.",
    "**Agregační pipeline** – 90+ operátorů pro in-database analytiku.",
    "**Automatický failover** – Raft-based election out-of-the-box.",
    "**ACID transakce** (od 4.0) – multi-document transactions.",
]:
    bullet(v)

h2("Obecné nevýhody MongoDB")
for n in [
    "**RAM náročnost** – WiredTiger cache 50 % RAM, malé nody jsou snadno přetíženy.",
    "**16 MB limit dokumentu** – nelze uložit velké binární soubory bez GridFS.",
    "**Sekundární indexy zpomalují zápisy** – každý index aktualizovaný při insert/update.",
    "**$lookup cross-shard** – JOIN přes shardované kolekce je neefektivní.",
    "**SSPL licence** – od roku 2018 omezuje cloud re-distribution.",
]:
    bullet(n)

h2("Výhody tohoto řešení")
for v in [
    "Plná automatizace přes `docker compose up -d` – žádný manuální zásah.",
    "Dynamický keyfile – splňuje požadavek zadání, nikdy není statický.",
    "Two-phase inicializace – elegantně řeší chicken-and-egg problém uživatelů s keyfile.",
    "Kompletní `$jsonSchema` validační schémata pro všechny 3 kolekce.",
    "30 netriviálních dotazů demonstrujících širokou paletu MongoDB schopností.",
]:
    bullet(v)

h2("Nevýhody tohoto řešení")
for n in [
    "Single-host deployment – fyzická fault tolerance je iluzorní (pád hostitele = pád clusteru).",
    "Hesla v `docker-compose.yml` – v produkci by byla v Docker Secrets nebo externím vault.",
    "Chybí TLS mezi kontejnery – bezpečné pouze v izolované Docker network.",
    "Dataset `ratings_small.csv` (100k) je relativně malý pro demonstraci shardingu.",
]:
    bullet(n)

# ══════════════════════════════════════════════════════════════════════════
# 5. DATA
# ══════════════════════════════════════════════════════════════════════════
h1("5. Data")

h2("5.1 Zdroje dat")
table(
    ["Dataset", "Zdroj", "Záznamy"],
    [
        ["TMDB 5000 Movies", "kaggle.com/datasets/tmdb/tmdb-movie-metadata", "4 803"],
        ["TMDB 5000 Credits", "kaggle.com/datasets/tmdb/tmdb-movie-metadata", "4 803"],
        ["MovieLens Ratings (Small)", "kaggle.com/datasets/rounakbanik/the-movies-dataset", "100 004"],
    ]
)
para("**Vazby:** `movies.id ↔ credits.movie_id` (1:1), `movies.id ↔ ratings.movieId` (1:N). Celkem **109 610 dokumentů** v databázi `filmdb`.")

h2("5.2 ETL pipeline (import_data.py)")
bullet("Načtení CSV přes `pd.read_csv`, typové konverze (`int`, `double`).")
bullet("Parse stringified JSON sloupců (`genres`, `cast`, `crew`) přes `ast.literal_eval` – TMDB CSV používá single quotes.")
bullet("Fill NA: čísla → 0, stringy → \"\". Drop řádků s chybějícím `id` / `movie_id`.")
bullet("Import v batchích po 1 000 záznamech přes `insert_many()` – 100× rychlejší než `insert_one()`.")

h2("5.3 EDA (eda_import.ipynb)")
bullet("Průměrný rozpočet: $29M, průměrné tržby: $82M → průměrný ROI ~3×.")
bullet("Top žánry: Drama, Comedy, Thriller, Action, Romance.")
bullet("671 unikátních uživatelů, 9 066 unikátních filmů; průměrné hodnocení 3.54 hvězd.")
bullet("Power-law distribuce hodnocení na uživatele – nejaktivnější uživatelé jsou outliery (1 000+ hodnocení).")
bullet("Nárůst filmové produkce od 80. let – exponenciální trend v datasetu.")

h2("5.4 Validační schémata ($jsonSchema)")
para("Vytvořena přes `collMod` (kolekce existují po `sh.shardCollection()`) v `05-init-validation.sh`. Každá kolekce má definovaná povinná pole, typy a rozsahy hodnot.")
table(
    ["Kolekce", "Povinná pole", "Příklady validací"],
    [
        ["movies", "id, title, vote_average", "vote_average: 0–10, budget ≥ 0"],
        ["credits", "movie_id, cast, crew", "cast a crew jsou array"],
        ["ratings", "userId, movieId, rating", "rating: 0.5–5.0"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════
# 6. DOTAZY
# ══════════════════════════════════════════════════════════════════════════
h1("6. Dotazy")
para("Kompletní syntax všech dotazů v souboru `dotazy/dotazy.md`. Připojení: `docker exec -it mongos1 mongosh -u admin -p adminpass123 --authenticationDatabase admin`, pak `use filmdb`.")
table(
    ["Kategorie", "Dotazy", "Klíčové techniky"],
    [
        ["1. Agregační a analytické", "1–6", "$unwind, $group, $avg, $bucket, $facet, $sum"],
        ["2. Propojování dat", "7–12", "$lookup, $filter, $arrayElemAt, $stdDevSamp"],
        ["3. Transformace dat", "13–18", "$addFields, $switch, $cond, $unset, $map, $toDate"],
        ["4. Indexy a optimalizace", "19–24", "explain(), hint(), $indexStats, $collStats"],
        ["5. Distribuce dat a cluster", "25–30", "sh.status(), rs.status(), getShardDistribution()"],
    ]
)

# ── KATEGORIE 1 ──────────────────────────────────────────────────────────
h2("Kategorie 1: Agregační a analytické dotazy")

query(1, "Průměrné hodnocení filmů podle žánru",
"Průměrné hodnocení a počet filmů na žánr (min. 10 filmů), seřazeno od nejlepšího.",
"""db.movies.aggregate([
  { $unwind: "$genres" },
  { $group: { _id: "$genres.name",
      prumerne_hodnoceni: { $avg: "$vote_average" },
      pocet_filmu: { $sum: 1 } } },
  { $match: { pocet_filmu: { $gte: 10 } } },
  { $sort: { prumerne_hodnoceni: -1 } },
  { $project: { _id: 1, pocet_filmu: 1,
      prumerne_hodnoceni: { $round: ["$prumerne_hodnoceni", 2] } } }
])""",
"`$unwind` rozloží pole `genres` – každý žánr se stane samostatným dokumentem. `$match` po `$group` filtruje žánry s malou vzorkou. Scatter-gather přes všechny shardy (shard key `id` není v filtru).")

query(2, "Top 10 nejziskovějších filmů s výpočtem ROI",
"Filmy s nejvyšším absolutním ziskem a ROI (%), filtrováno na reálná finanční data.",
"""db.movies.aggregate([
  { $match: { budget: { $gt: 1000000 }, revenue: { $gt: 0 } } },
  { $addFields: {
      zisk: { $subtract: ["$revenue", "$budget"] },
      roi_procent: { $round: [{ $multiply: [
          { $divide: [{ $subtract: ["$revenue","$budget"] }, "$budget"] }, 100
      ] }, 1] } } },
  { $sort: { zisk: -1 } },
  { $limit: 10 },
  { $project: { _id:0, title:1,
      budget_mil: { $round: [{ $divide: ["$budget", 1000000] }, 1] },
      zisk_mil:   { $round: [{ $divide: ["$zisk", 1000000] }, 1] },
      roi_procent: 1, vote_average: 1 } }
])""",
"`$addFields` přidá `zisk` a `roi_procent` jako odvozená pole v jediném průchodu. `$project` transformuje USD na miliony pro čitelnost.")

query(3, "Distribuce filmů do hodnotících skupin pomocí $bucket",
"Rozdělit filmy do skupin (0–4, 4–5, 5–6, 6–7, 7–8, 8–10) s průměrným rozpočtem a tržbami.",
"""db.movies.aggregate([
  { $match: { vote_count: { $gte: 50 } } },
  { $bucket: {
      groupBy: "$vote_average",
      boundaries: [0, 4, 5, 6, 7, 8, 10.1],
      default: "Nehodnoceno",
      output: {
        pocet_filmu: { $sum: 1 },
        prumerny_budget_mil: { $avg: { $divide: ["$budget", 1000000] } },
        prumerne_trzby_mil:  { $avg: { $divide: ["$revenue", 1000000] } } } } },
  { $project: { _id:1, pocet_filmu:1,
      prumerny_budget_mil: { $round: ["$prumerny_budget_mil", 1] },
      prumerne_trzby_mil:  { $round: ["$prumerne_trzby_mil", 1] } } }
])""",
"`$bucket` rozdělí dokumenty do předdefinovaných intervalů. `default` zachytí hodnoty mimo rozsah. Výsledek odhalí korelaci hodnocení s rozpočty.")

query(4, "Roční analýza filmové produkce (2000–2016)",
"Počet filmů, průměrné hodnocení, celkový rozpočet a tržby pro každý rok.",
"""db.movies.aggregate([
  { $addFields: { rok: { $toInt: { $substr: ["$release_date", 0, 4] } } } },
  { $match: { rok: { $gte: 2000, $lte: 2016 }, vote_count: { $gte: 20 } } },
  { $group: { _id: "$rok",
      pocet_filmu: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$vote_average" },
      celkovy_budget_mld: { $sum: "$budget" },
      celkove_trzby_mld:  { $sum: "$revenue" } } },
  { $sort: { _id: 1 } },
  { $project: { rok: "$_id", pocet_filmu: 1,
      prumerne_hodnoceni: { $round: ["$prumerne_hodnoceni", 2] },
      celkovy_budget_mld: { $round: [{ $divide: ["$celkovy_budget_mld", 1e9] }, 2] },
      celkove_trzby_mld:  { $round: [{ $divide: ["$celkove_trzby_mld", 1e9] }, 2] } } }
])""",
"`$substr` + `$toInt` extrahuje rok z řetězce `release_date`. Více akumulátorů (`$avg`, `$sum`, `$max`) v jednom `$group`. `$project` přejmenuje `_id` na `rok`.")

query(5, "Filmové produkční společnosti s nejlepším výkonem",
"Společnosti s min. 5 filmy, seřazené podle průměrného hodnocení s celkovým ROI.",
"""db.movies.aggregate([
  { $unwind: "$production_companies" },
  { $group: { _id: "$production_companies.name",
      pocet_filmu: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$vote_average" },
      celkove_trzby: { $sum: "$revenue" },
      celkovy_budget: { $sum: "$budget" } } },
  { $match: { pocet_filmu: { $gte: 5 }, celkovy_budget: { $gt: 0 } } },
  { $addFields: { efektivita_roi: { $round: [{ $multiply: [
      { $divide: [{ $subtract: ["$celkove_trzby","$celkovy_budget"] }, "$celkovy_budget"] },
      100 ] }, 1] } } },
  { $sort: { prumerne_hodnoceni: -1 } },
  { $limit: 10 }
])""",
"`$match` po `$group` filtruje malé filmografie a nulové budgety (zamezení dělení nulou). `$addFields` přidá ROI metriku po agregaci.")

query(6, "Multi-dimenzionální analýza pomocí $facet",
"Paralelní analýza ze tří pohledů: top žánry, hodnotící skupiny a produkce po dekádách.",
"""db.movies.aggregate([
  { $match: { vote_count: { $gte: 100 }, release_date: { $gt: "" } } },
  { $facet: {
    "top_zanry": [
      { $unwind: "$genres" },
      { $group: { _id: "$genres.name", pocet: { $sum: 1 }, avg_hodnoceni: { $avg: "$vote_average" } } },
      { $sort: { pocet: -1 } }, { $limit: 5 }
    ],
    "hodnotici_skupiny": [
      { $bucket: { groupBy: "$vote_average",
          boundaries: [0,5,6,7,8,10.1], default: "N/A",
          output: { pocet: { $sum: 1 } } } }
    ],
    "po_dekadach": [
      { $addFields: { dekada: { $multiply: [
          { $floor: { $divide: [{ $toInt: { $substr: ["$release_date",0,4] } }, 10] } }, 10 ] } } },
      { $group: { _id: "$dekada", pocet: { $sum: 1 }, avg_hodnoceni: { $avg: "$vote_average" } } },
      { $sort: { _id: 1 } }
    ] } }
])""",
"`$facet` spustí nezávislé pipeline nad stejnými daty v jediném průchodu kolekcí – efektivnější než 3 separátní dotazy. Vrátí jeden dokument se třemi vnořenými výsledky.")

# ── KATEGORIE 2 ──────────────────────────────────────────────────────────
h2("Kategorie 2: Propojování dat a vazby")

query(7, "Nejčastěji obsazovaní herci",
"Herci s největším počtem filmů v databázi credits, s ukázkou filmografie.",
"""db.credits.aggregate([
  { $unwind: "$cast" },
  { $group: { _id: "$cast.name",
      pocet_filmu: { $sum: 1 }, filmy: { $push: "$title" } } },
  { $sort: { pocet_filmu: -1 } },
  { $limit: 10 },
  { $project: { _id:1, pocet_filmu:1,
      ukazka_filmu: { $slice: ["$filmy", 3] } } }
])""",
"`$unwind` rozloží `cast` – z 4 800 creditů vzniknou stovky tisíc dokumentů. `$slice` zobrazí jen 3 filmy jako ukázku.")

query(8, "Nejlépe hodnocené filmy s jejich režiséry",
"Ke každému vysoce hodnocenému filmu (≥ 7.5, min. 500 hlasů) dohledat režiséra z credits.",
"""db.movies.aggregate([
  { $match: { vote_average: { $gte: 7.5 }, vote_count: { $gte: 500 } } },
  { $lookup: { from:"credits", localField:"id", foreignField:"movie_id", as:"crew_info" } },
  { $unwind: { path:"$crew_info", preserveNullAndEmptyArrays: false } },
  { $addFields: { reziser: { $arrayElemAt: [
      { $filter: { input:"$crew_info.crew", as:"c",
          cond: { $eq: ["$$c.job","Director"] } } }, 0 ] } } },
  { $project: { _id:0, title:1, vote_average:1, release_date:1,
      reziser_jmeno: "$reziser.name" } },
  { $sort: { vote_average: -1 } }, { $limit: 15 }
])""",
"`$lookup` realizuje JOIN `movies.id → credits.movie_id`. `$filter` vrátí členy crew s `job = Director`. `$arrayElemAt` vezme prvního nalezeného.")

query(9, "Nejúspěšnější režiséři podle průměrného hodnocení",
"Režiséři s min. 3 filmy a min. 100 hlasů na film, seřazeni podle průměrného hodnocení filmografie.",
"""db.credits.aggregate([
  { $unwind: "$crew" },
  { $match: { "crew.job": "Director" } },
  { $lookup: { from:"movies", localField:"movie_id", foreignField:"id", as:"film_info" } },
  { $unwind: "$film_info" },
  { $match: { "film_info.vote_count": { $gte: 100 } } },
  { $group: { _id: "$crew.name",
      pocet_filmu: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$film_info.vote_average" },
      filmy: { $push: "$film_info.title" } } },
  { $match: { pocet_filmu: { $gte: 3 } } },
  { $sort: { prumerne_hodnoceni: -1 } }, { $limit: 10 }
])""",
"Pipeline začíná v `credits`, filtruje Directors před `$lookup` – sníží objem dat. Dva `$match` v různých fázích: vstup i výstup `$lookup`.")

query(10, "Herci s nejvíce hlavními rolemi v akčních filmech",
"Herci (role order ≤ 3) v akčních filmech, seřazeni podle počtu hlavních rolí.",
"""db.movies.aggregate([
  { $unwind: "$genres" }, { $match: { "genres.name": "Action" } },
  { $group: { _id: "$id", title: { $first: "$title" }, vote_average: { $first: "$vote_average" } } },
  { $lookup: { from:"credits", localField:"_id", foreignField:"movie_id", as:"cast_info" } },
  { $unwind: "$cast_info" }, { $unwind: "$cast_info.cast" },
  { $match: { "cast_info.cast.order": { $lte: 3 } } },
  { $group: { _id: "$cast_info.cast.name",
      pocet_action_filmu: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$vote_average" } } },
  { $sort: { pocet_action_filmu: -1 } }, { $limit: 10 }
])""",
"7 fází pipeline: `$unwind` + `$match` pro akční filmy, `$group` odstraní duplicity, `$lookup` přidá herce, 2 `$unwind` rozloží vnořená data, `$match` omezí na protagonisty.")

query(11, "Vliv velikosti hereckého obsazení na hodnocení a tržby",
"Korelace počtu herců (kategorie) s průměrným hodnocením a tržbami.",
"""db.movies.aggregate([
  { $match: { vote_count: { $gte: 100 } } },
  { $lookup: { from:"credits", localField:"id", foreignField:"movie_id", as:"credits_data" } },
  { $unwind: { path:"$credits_data", preserveNullAndEmptyArrays: false } },
  { $addFields: { pocet_hercu: { $size: "$credits_data.cast" } } },
  { $group: { _id: { $switch: { branches: [
      { case: { $lte: ["$pocet_hercu", 10] }, then: "01–10 herců" },
      { case: { $lte: ["$pocet_hercu", 20] }, then: "11–20 herců" },
      { case: { $lte: ["$pocet_hercu", 50] }, then: "21–50 herců" }
    ], default: "50+ herců" } },
    pocet_filmu: { $sum: 1 },
    prumerne_hodnoceni: { $avg: "$vote_average" },
    prumerne_trzby_mil: { $avg: { $divide: ["$revenue", 1000000] } } } },
  { $sort: { _id: 1 } }
])""",
"`$size` spočítá počet herců z vnořeného pole přineseného `$lookup`. `$switch` přímo v `$group._id` jako skupinový klíč – netriviální použití výrazu.")

query(12, "Nejhodnocenější filmy uživateli s analýzou konzistence",
"Filmy s největším počtem hodnocení a standardní odchylka (konzistentní vs kontroverzní).",
"""db.ratings.aggregate([
  { $group: { _id: "$movieId",
      pocet_hodnoceni: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$rating" },
      smerodatna_odchylka: { $stdDevSamp: "$rating" },
      unikatni_hodnotitele: { $addToSet: "$userId" } } },
  { $addFields: {
      pocet_unikatnich: { $size: "$unikatni_hodnotitele" },
      konzistence: { $switch: { branches: [
        { case: { $lt: ["$smerodatna_odchylka", 0.8] }, then: "Konzistentní" },
        { case: { $lt: ["$smerodatna_odchylka", 1.2] }, then: "Středně variabilní" }
      ], default: "Kontroverzní" } } } },
  { $match: { pocet_hodnoceni: { $gte: 50 } } },
  { $sort: { pocet_hodnoceni: -1 } }, { $limit: 10 }
])""",
"`$stdDevSamp` je statistický akumulátor pro výběrovou odchylku. `$addToSet` shromáždí unikátní userId, `$size` spočítá jejich počet po `$group`.")

# ── KATEGORIE 3 ──────────────────────────────────────────────────────────
h2("Kategorie 3: Transformace a obohacení dat")

query(13, "Klasifikace filmů podle rozpočtu a analýza ROI",
"Kategorie blockbuster/velká/střední/nízkorozpočtová s průměrným ROI a procentem ziskových filmů.",
"""db.movies.aggregate([
  { $match: { budget: { $gt: 0 }, revenue: { $gt: 0 } } },
  { $addFields: {
      kategorie: { $switch: { branches: [
        { case: { $gte: ["$budget", 100000000] }, then: "Blockbuster (100M+)" },
        { case: { $gte: ["$budget", 20000000] }, then: "Velká produkce (20–100M)" },
        { case: { $gte: ["$budget", 5000000] }, then: "Středorozpočtový (5–20M)" }
      ], default: "Nízkorozpočtový (<5M)" } },
      roi: { $round: [{ $multiply: [
        { $divide: [{ $subtract: ["$revenue","$budget"] }, "$budget"] }, 100 ] }, 1] } } },
  { $group: { _id: "$kategorie",
      pocet_filmu: { $sum: 1 },
      prumerne_roi: { $avg: "$roi" },
      prumerne_hodnoceni: { $avg: "$vote_average" },
      uspesnych_podil: { $avg: { $cond: [{ $gt: ["$revenue","$budget"] }, 1, 0] } } } },
  { $sort: { prumerne_roi: -1 } }
])""",
"`$avg` na podmíněném výrazu `$cond` vypočítá procento ziskových filmů bez separátního `$match`. Ukazuje, že nízkorozpočtové filmy mohou mít vyšší ROI.")

query(14, "Analýza mluvených jazyků ve světové kinematografii",
"Rozložení jazyků, průměrné hodnocení a tržby pro každý jazyk (min. 15 filmů).",
"""db.movies.aggregate([
  { $unwind: "$spoken_languages" },
  { $group: { _id: "$spoken_languages.name",
      pocet_filmu: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$vote_average" },
      celkove_trzby: { $sum: "$revenue" } } },
  { $match: { pocet_filmu: { $gte: 15 } } },
  { $addFields: { podil_trzby_mld: { $round: [{ $divide: ["$celkove_trzby",1e9] }, 2] } } },
  { $sort: { pocet_filmu: -1 } }, { $limit: 12 }
])""",
"`$unwind` rozloží `spoken_languages` – film může být ve více jazycích. `$addFields` po agregaci přidá celkový tržní podíl v miliardách.")

query(15, "Obohacení dokumentu – přehledná karta filmu",
"Přidat odvozená pole (rok, normalizované hodnocení, délka, finanční úspěch, seznam žánrů) a odebrat přebytečná pole.",
"""db.movies.aggregate([
  { $match: { vote_count: { $gte: 200 }, release_date: { $gt: "" } } },
  { $addFields: {
      rok_vydani: { $toInt: { $substr: ["$release_date", 0, 4] } },
      hodnoceni_100: { $round: [{ $multiply: ["$vote_average", 10] }, 0] },
      klasifikace_delky: { $switch: { branches: [
        { case: { $lt: ["$runtime", 90] }, then: "Krátký" },
        { case: { $lte: ["$runtime", 120] }, then: "Standardní" },
        { case: { $lte: ["$runtime", 150] }, then: "Dlouhý" }
      ], default: "Epický" } },
      financni_uspech: { $cond: { if: { $and: [{ $gt: ["$budget",0] }, { $gt: ["$revenue",0] }] },
        then: { $cond: [{ $gt: ["$revenue", { $multiply: ["$budget",2] }] }, "Velký úspěch",
          { $cond: [{ $gt: ["$revenue","$budget"] }, "Ziskový", "Ztrátový"] }] },
        else: "Data nedostupná" } },
      zanry_seznam: { $map: { input:"$genres", as:"g", in:"$$g.name" } } } },
  { $unset: ["genres","keywords","production_companies","production_countries",
             "spoken_languages","homepage","tagline","popularity","_id"] },
  { $sort: { vote_average: -1 } }, { $limit: 10 }
])""",
"`$addFields` + `$unset` je standardní vzor pro transformaci dokumentů. `$map` transformuje pole objektů na pole řetězců. Vnořené `$cond` realizuje trojstavovou logiku.")

query(16, "Trendy filmového průmyslu po pětiletích (1980–2016)",
"Vývoj průměrného hodnocení, délky filmů, objemu produkce a investic v 5-letých intervalech.",
"""db.movies.aggregate([
  { $addFields: { rok: { $toInt: { $substr: ["$release_date",0,4] } } } },
  { $match: { rok: { $gte: 1980, $lte: 2016 }, vote_count: { $gte: 30 } } },
  { $addFields: { petileti: { $multiply: [{ $floor: { $divide: ["$rok",5] } }, 5] } } },
  { $group: { _id: "$petileti",
      pocet_filmu: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$vote_average" },
      prumerny_runtime: { $avg: "$runtime" },
      prumerny_budget_mil: { $avg: { $divide: ["$budget",1e6] } } } },
  { $sort: { _id: 1 } },
  { $project: { obdobi: { $concat: [{ $toString: "$_id" }, "–",
      { $toString: { $add: ["$_id",4] } }] },
    pocet_filmu:1, prumerne_hodnoceni:{ $round: ["$prumerne_hodnoceni",2] },
    prumerny_runtime:{ $round: ["$prumerny_runtime",0] },
    prumerny_budget_mil:{ $round: ["$prumerny_budget_mil",1] } } }
])""",
"Dvoustupňový `$addFields`: extrakce roku, pak výpočet pětiletí jako `floor(rok/5)*5`. `$concat` + `$toString` + `$add` vytvoří čitelný label \"1980–1984\".")

query(17, "Analýza aktivity uživatelů v čase",
"Počty hodnocení a průměrné hodnocení po měsících (Unix timestamp → Date).",
"""db.ratings.aggregate([
  { $addFields: { datum: { $toDate: { $multiply: ["$timestamp", 1000] } } } },
  { $addFields: { rok: { $year: "$datum" }, mesic: { $month: "$datum" } } },
  { $group: { _id: { rok: "$rok", mesic: "$mesic" },
      pocet_hodnoceni: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$rating" },
      unikatni_uzivatele: { $addToSet: "$userId" } } },
  { $addFields: { pocet_uzivatelu: { $size: "$unikatni_uzivatele" } } },
  { $sort: { "_id.rok": 1, "_id.mesic": 1 } },
  { $project: { _id:0, rok:"$_id.rok", mesic:"$_id.mesic",
      pocet_hodnoceni:1, prumerne_hodnoceni:{ $round:["$prumerne_hodnoceni",2] },
      pocet_uzivatelu:1 } },
  { $limit: 24 }
])""",
"`$multiply` převede Unix timestamp (sec) na ms, `$toDate` vytvoří Date objekt. `$year` a `$month` extrahují části – dostupné pouze na Date typech. Složený klíč v `$group._id`.")

query(18, "Segmentace uživatelů podle hodnotícího chování",
"Uživatelé segmentováni do skupin (aktivita × styl: optimista/kritik/vyrovnaný).",
"""db.ratings.aggregate([
  { $group: { _id: "$userId",
      pocet_hodnoceni: { $sum: 1 },
      prumerne_hodnoceni: { $avg: "$rating" } } },
  { $addFields: {
      typ_uzivatele: { $switch: { branches: [
        { case: { $gte: ["$pocet_hodnoceni", 200] }, then: "Superhodnotitel (200+)" },
        { case: { $gte: ["$pocet_hodnoceni", 50] },  then: "Aktivní (50–199)" },
        { case: { $gte: ["$pocet_hodnoceni", 20] },  then: "Příležitostný (20–49)" }
      ], default: "Nový (<20)" } },
      styl_hodnoceni: { $switch: { branches: [
        { case: { $gt: ["$prumerne_hodnoceni", 4.0] }, then: "Optimista" },
        { case: { $lt: ["$prumerne_hodnoceni", 2.5] }, then: "Kritik" }
      ], default: "Vyrovnaný" } } } },
  { $group: { _id: { typ: "$typ_uzivatele", styl: "$styl_hodnoceni" },
      pocet_uzivatelu: { $sum: 1 },
      prumerny_pocet_hodnoceni: { $avg: "$pocet_hodnoceni" } } },
  { $sort: { pocet_uzivatelu: -1 } }
])""",
"Dvoustupňová agregace: první `$group` agreguje za uživatele, `$addFields` klasifikuje, druhý `$group` seskupí do segmentů (typ × styl).")

# ── KATEGORIE 4 ──────────────────────────────────────────────────────────
h2("Kategorie 4: Indexy a optimalizace")

query(19, "Srovnání výkonu dotazu s indexem a bez indexu",
"explain(executionStats) na neindexované vs indexované pole – COLLSCAN vs IXSCAN.",
"""// Bez indexu → COLLSCAN (pole original_language neindexováno)
db.movies.aggregate([
  { $match: { original_language: "fr", vote_count: { $gte: 100 } } },
  { $group: { _id: "$original_language",
      pocet_filmu: { $sum: 1 }, prumerne_hodnoceni: { $avg: "$vote_average" } } }
]).explain("executionStats")

// S indexem → IXSCAN (compound index { vote_average: -1, vote_count: -1 })
db.movies.aggregate([
  { $match: { vote_average: { $gte: 8.0 }, vote_count: { $gte: 200 } } },
  { $group: { _id: { $floor: "$vote_average" },
      pocet_filmu: { $sum: 1 },
      prumerny_budget_mil: { $avg: { $divide: ["$budget", 1e6] } } } },
  { $sort: { _id: -1 } }
]).explain("executionStats")""",
"Bez indexu uvidíme `stage: COLLSCAN` a `docsExamined = ~4 800`. S indexem uvidíme `IXSCAN` a výrazně nižší `totalKeysExamined`. V shardovaném clusteru `explain()` odhalí navíc `SHARD_MERGE`.")

query(20, "Fulltextové vyhledávání s relevance score",
"Vyhledat filmy s tématikou vesmíru pomocí textového indexu a seřadit podle relevance.",
"""db.movies.aggregate([
  { $match: { $text: { $search: "space alien future war galaxy" } } },
  { $addFields: { skore_relevance: { $meta: "textScore" } } },
  { $match: { vote_count: { $gte: 100 } } },
  { $sort: { skore_relevance: -1 } },
  { $project: { _id:0, title:1, vote_average:1,
      skore_relevance: { $round: ["$skore_relevance", 3] },
      ukazka_popisu: { $substr: ["$overview", 0, 120] } } },
  { $limit: 10 }
])""",
"`$text` využívá textový index `{ title: \"text\", overview: \"text\" }` – bez indexu by dotaz selhal. `$meta: \"textScore\"` přidá relevance skóre. Index používá stemming a ignoruje stopwords.")

query(21, "Vynucení konkrétního indexu pomocí hint()",
"Porovnat plán při automatickém výběru indexu, vynuceném compound indexu a COLLSCAN.",
"""// Varianta A: automatický výběr optimizerem
db.movies.aggregate([
  { $match: { vote_average: { $gte: 7.0 }, vote_count: { $gte: 100 } } },
  { $group: { _id: null, pocet_filmu: { $sum: 1 }, prumerne_hodnoceni: { $avg: "$vote_average" } } }
]).explain("executionStats")

// Varianta B: explicitní compound index
db.movies.aggregate(
  [{ $match: { vote_average: { $gte: 7.0 }, vote_count: { $gte: 100 } } },
   { $group: { _id: null, pocet_filmu: { $sum: 1 } } }],
  { hint: { vote_average: -1, vote_count: -1 } }
).explain("executionStats")

// Varianta C: vynucení COLLSCAN
db.movies.aggregate(
  [{ $match: { vote_average: { $gte: 7.0 }, vote_count: { $gte: 100 } } },
   { $group: { _id: null, pocet_filmu: { $sum: 1 } } }],
  { hint: "$natural" }
).explain("executionStats")""",
"`hint()` jako volba agregační pipeline (MongoDB 4.2+) vynutí konkrétní index nebo COLLSCAN. `$natural` vynutí COLLSCAN i přes existující indexy – `executionTimeMillis` výrazně vzroste.")

query(22, "Statistiky využití indexů v kolekci movies",
"Zobrazit všechny indexy, jejich počet použití a identifikovat případné nevyužívané.",
"""db.movies.aggregate([
  { $indexStats: {} },
  { $project: {
      nazev_indexu: "$name",
      pocet_pouziti: "$accesses.ops",
      sledovano_od: "$accesses.since",
      klic_indexu: "$key" } },
  { $sort: { pocet_pouziti: -1 } }
])""",
"`$indexStats` vrátí statistiky každého indexu bez dalšího filtru. V shardovaném clusteru agreguje statistiky ze všech shardů. Indexy s `pocet_pouziti: 0` jsou nevyužívané a zbytečně zpomalují zápisy.")

query(23, "Analytická kontrola referenční integrity dat",
"Osiřelá hodnocení (bez záznamu v movies), osiřelé credits a filmy s anomálním ROI.",
"""// Kontrola 1: Hodnocení bez odpovídajícího záznamu v movies
db.ratings.aggregate([
  { $group: { _id: "$movieId", pocet_hodnoceni: { $sum: 1 } } },
  { $lookup: { from:"movies", localField:"_id", foreignField:"id", as:"film_data" } },
  { $match: { film_data: { $size: 0 } } },
  { $sort: { pocet_hodnoceni: -1 } }, { $limit: 10 }
])

// Kontrola 2: Credits bez záznamu v movies
db.credits.aggregate([
  { $lookup: { from:"movies", localField:"movie_id", foreignField:"id", as:"film_data" } },
  { $match: { film_data: { $size: 0 } } },
  { $project: { _id:0, movie_id:1, title:1 } }, { $limit: 10 }
])

// Kontrola 3: Filmy s podezřele vysokým ROI (revenue > 50× budget)
db.movies.aggregate([
  { $match: { budget: { $gt: 1000000 }, revenue: { $gt: 0 } } },
  { $addFields: { roi_nasobek: { $round: [{ $divide: ["$revenue","$budget"] }, 1] } } },
  { $match: { roi_nasobek: { $gt: 50 } } },
  { $sort: { roi_nasobek: -1 } }, { $limit: 10 }
])""",
"Tři kontroly: osiřelá data (film_data prázdné po `$lookup`), anomální finanční data. MovieLens obsahuje filmy neexistující v TMDB – dokumentuje kvalitu datasetů.")

query(24, "Targeted query vs scatter-gather (ratings)",
"Porovnat efektivitu dotazu přes shard key (1 shard) oproti dotazu bez shard key (3 shardy).",
"""// Targeted query – přes shard key userId → 1 shard (SHARDS_SCANNED: 1)
db.ratings.find({ userId: 42 }).explain("executionStats")

// Scatter-gather – bez shard key → všechny 3 shardy (SHARDS_SCANNED: 3)
db.ratings.find(
  { rating: { $gte: 4.5 }, movieId: { $in: [1, 2, 50, 110, 260] } }
).explain("executionStats")

// Analytický dotaz s sekundárním indexem na movieId
db.ratings.aggregate([
  { $match: { movieId: 296 } },
  { $group: { _id: "$rating", pocet: { $sum: 1 } } },
  { $sort: { _id: -1 } }
])""",
"Targeted query s `userId` (shard key) jde na 1 shard. Scatter-gather bez shard key prochází všechny 3 shardy paralelně. Sekundární index na `movieId` pomáhá v rámci každého shardu.")

# ── KATEGORIE 5 ──────────────────────────────────────────────────────────
h2("Kategorie 5: Distribuce dat, cluster a replikace")

query(25, "Kompletní stav shardovaného clusteru",
"sh.status(), stav balanceru a přehled aktivních mongos routerů z config databáze.",
"""// Spustit na mongos1
sh.status()

// Stav balanceru a migrace chunků
db.adminCommand({ balancerStatus: 1 })

// Přehled aktivních mongos routerů z config DB
use config
db.mongos.aggregate([
  { $addFields: { posledni_ping: {
      $toDate: { $multiply: [{ $tsSecond: "$ping" }, 1000] } } } },
  { $project: { _id:0, adresa:"$_id", verze:"$mongoVersion", posledni_ping:1 } },
  { $sort: { adresa: 1 } }
])""",
"`sh.status()` zobrazí shardy, chunky, balancer. `balancerStatus` odhalí probíhající migrace. Pipeline nad `config.mongos` ukáže čas posledního heartbeatu každého routeru – mongos bez pingu je pravděpodobně nedostupný.")

query(26, "Distribuce dat na shardech pro všechny kolekce",
"getShardDistribution() pro movies, credits, ratings + $collStats pro počty dokumentů.",
"""use filmdb
db.movies.getShardDistribution()
db.credits.getShardDistribution()
db.ratings.getShardDistribution()

// Přesné počty dokumentů a velikost na každém shardu (ratings)
db.ratings.aggregate([
  { $collStats: { storageStats: {} } },
  { $project: { _id:0, shard:1,
      pocet_dokumentu: "$storageStats.count",
      velikost_mb: { $round: [{ $divide: ["$storageStats.storageSize", 1048576] }, 2] },
      prumerny_dokument_b: { $round: ["$storageStats.avgObjSize", 0] },
      pocet_indexu: "$storageStats.nindexes" } },
  { $sort: { pocet_dokumentu: -1 } }
])""",
"`getShardDistribution()` zobrazí % dat na každém shardu. `$collStats` s `storageStats` vrátí per-shard statistiky ze storage enginu. Hashed sharding zaručuje odchylku < 5 % mezi shardy.")

query(27, "Stav replica setu a detaily replikace shard1",
"rs.status(), rs.conf() a analýza oplog logu za posledních 24 hodin.",
"""// docker exec -it shard1svr1 mongosh -u admin -p adminpass123 --authenticationDatabase admin
rs.status()
rs.conf()

// Typy operací v oplog logu PRIMARY za 24 hodin
use local
db.oplog.rs.aggregate([
  { $match: { ts: { $gte: new Timestamp(Math.floor(Date.now()/1000) - 86400, 0) } } },
  { $group: { _id: "$op", pocet: { $sum: 1 } } },
  { $addFields: { popis: { $switch: { branches: [
      { case: { $eq: ["$_id","i"] }, then: "INSERT" },
      { case: { $eq: ["$_id","u"] }, then: "UPDATE" },
      { case: { $eq: ["$_id","d"] }, then: "DELETE" },
      { case: { $eq: ["$_id","n"] }, then: "NOOP (heartbeat)" }
    ], default: "COMMAND" } } } },
  { $sort: { pocet: -1 } }
])""",
"`rs.status()` zobrazí `stateStr`, `health`, `optimeDate` a `pingMs`. `rs.conf()` doplní `heartbeatIntervalMillis` (2000 ms) a `electionTimeoutMillis` (10 000 ms). Oplog pipeline ukáže typy replikovaných operací.")

query(28, "Simulace výpadku PRIMARY nodu a sledování election",
"Zastavit PRIMARY shard1svr1, sledovat Raft election a ověřit dostupnost clusteru.",
"""// KROK 1: Zjistit aktuálního PRIMARY
rs.isMaster()""",
"")
code("""# KROK 2: Zastavit PRIMARY kontejner (simulace výpadku)
docker stop shard1svr1

# KROK 3: Připojit se na shard1svr2
docker exec -it shard1svr2 mongosh -u admin -p adminpass123 --authenticationDatabase admin""")
code("""// KROK 4: Sledovat průběh election (opakovat ~30 sekund)
rs.status().members.map(m => ({ name: m.name, state: m.stateStr, health: m.health }))""")
code("""# KROK 5: Obnovit výpadnutý node
docker start shard1svr1

// KROK 6: Ověřit znovupřipojení jako SECONDARY
rs.status().members.map(m => ({ name: m.name, state: m.stateStr }))""")
p = doc.add_paragraph()
_inline(p, "**Komentář:** Raft election probíhá 10–30 sekund. Zbývající 2 SECONDARY nody hlasují, nový PRIMARY je zvolen. Během election jsou zápisy dočasně nedostupné. Shard2 a shard3 zůstávají plně funkční – výpadek 1 shardu neovlivní ostatní.")
p.paragraph_format.space_after = Pt(5)

query(29, "Metadata shardů a distribuce chunků",
"listShards, shardované kolekce z config DB a počty chunků na shardech.",
"""// Seznam všech registrovaných shardů
db.adminCommand({ listShards: 1 })

// Metadata shardovaných kolekcí z config databáze
use config
db.collections.find({ _id: /^filmdb/ }, { _id:1, key:1, unique:1 })

// Počty chunků na shardech pro kolekce filmdb
db.chunks.aggregate([
  { $match: { ns: /^filmdb/ } },
  { $group: { _id: { ns:"$ns", shard:"$shard" }, pocet_chunku: { $sum: 1 } } },
  { $sort: { "_id.ns": 1, "_id.shard": 1 } }
])""",
"`config.collections` uchovává definice shardovaných kolekcí a shard keys. `config.chunks` mapuje chunky na shardy s hash rozsahy. Balancer přesouvá chunky, pokud rozdíl překročí threshold (výchozí 3 chunky).")

query(30, "Replikační lag a zdravotní stav replikace",
"Replikační zpoždění SECONDARY nodů, oplog okno a zdraví celé replikační skupiny.",
"""// docker exec -it shard1svr1 mongosh -u admin -p adminpass123 --authenticationDatabase admin
rs.printSecondaryReplicationInfo()
db.getReplicationInfo()

// Zdraví všech členů s lagy
rs.status().members.map(m => ({
  jmeno: m.name, stav: m.stateStr,
  zdravi: m.health, pingMs: m.pingMs || 0
}))

// Analytický výpočet oplog okna (jak daleko sahá replikační historie)
use local
db.oplog.rs.aggregate([
  { $group: { _id: null,
      nejstarsi_ts: { $min: "$ts" }, nejnovejsi_ts: { $max: "$ts" } } },
  { $project: { _id:0,
      nejstarsi_datum: { $toDate: { $multiply: [{ $tsSecond: "$nejstarsi_ts" }, 1000] } },
      nejnovejsi_datum: { $toDate: { $multiply: [{ $tsSecond: "$nejnovejsi_ts" }, 1000] } },
      oplog_okno_hodin: { $round: [{
          $divide: [{ $subtract: [{ $tsSecond: "$nejnovejsi_ts" }, { $tsSecond: "$nejstarsi_ts" }] },
          3600] }, 1] } } }
])""",
"`rs.printSecondaryReplicationInfo()` zobrazí lag za PRIMARY v sekundách (v naší konfiguraci < 1 s). `oplog_okno_hodin` je klíčová metrika – pokud byl SECONDARY offline déle než toto okno, musí provést full initial sync místo inkrementální replikace.")

# ══════════════════════════════════════════════════════════════════════════
# ZÁVĚR
# ══════════════════════════════════════════════════════════════════════════
h1("Závěr")
para("Projekt úspěšně demonstruje nasazení **MongoDB 8.0 Sharded Clusteru** s replikací, shardingem a zabezpečením prostřednictvím Docker Compose. Celá inicializace proběhne automaticky po jediném příkazu `docker compose up -d`.")
para("Klíčovými přínosy řešení jsou: dynamicky generovaný keyfile (bezpečná interní autentizace), dvoufázová inicializace s file-based synchronizací, kompletní validační schémata pro všechny tři kolekce a 30 netriviálních dotazů pokrývajících celé spektrum MongoDB schopností – od základní agregace přes cross-collection JOINy až po diagnostiku clusteru a simulaci výpadku.")
para("Filmová databáze (`filmdb`) s 109 610 dokumenty demonstruje, jak MongoDB řeší reálný problém kombinace strukturovaných dat (ratings) s vnořenými poli (cast, crew, genres) ve škálovatelné distribuované architektuře.")

# ══════════════════════════════════════════════════════════════════════════
# ZDROJE
# ══════════════════════════════════════════════════════════════════════════
h1("Zdroje a nástroje")
sources = [
    "Bradshaw, S., Brazil, E., Chodorow, K. *MongoDB: The Definitive Guide*. 3rd ed. O'Reilly Media, 2019.",
    "Brewer, E. A. \"Towards Robust Distributed Systems.\" PODC 2000.",
    "Docker Inc. *Docker Compose Documentation*. docs.docker.com, 2024.",
    "eBay Tech Blog. *Sharding MongoDB at eBay*. tech.ebay.com, 2012.",
    "Kaggle. *TMDB 5000 Movie Dataset*. kaggle.com/datasets/tmdb/tmdb-movie-metadata, 2017.",
    "Kaggle. *MovieLens Dataset*. kaggle.com/datasets/rounakbanik/the-movies-dataset, 2018.",
    "McKinney, W. *Python for Data Analysis*. 3rd ed. O'Reilly Media, 2022.",
    "MetLife. *Customer 360 with MongoDB*. mongodb.com/customers/metlife, 2014.",
    "MongoDB Inc. *MongoDB 8.0 Manual*. mongodb.com/docs/manual, 2024.",
    "MongoDB Inc. *MongoDB Sharding Introduction*. mongodb.com/docs/manual/sharding, 2024.",
    "The Weather Channel. *MongoDB Atlas Case Study*. mongodb.com/customers/the-weather-channel, 2015.",
]
for s in sorted(sources):
    bullet(s)

# ══════════════════════════════════════════════════════════════════════════
doc.save(OUT)
print(f"Hotovo: {OUT}")
