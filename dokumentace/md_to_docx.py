"""
Převede dokumentace.md na dokumentace.docx pomocí python-docx.
Spuštění: python dokumentace/md_to_docx.py
"""
import re
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

MD_FILE = os.path.join(os.path.dirname(__file__), "dokumentace.md")
OUT_FILE = os.path.join(os.path.dirname(__file__), "dokumentace.docx")


def set_heading_style(paragraph, level):
    sizes = {1: 22, 2: 18, 3: 14, 4: 13}
    bold  = {1: True, 2: True, 3: True, 4: True}
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.bold = bold.get(level, False)
    run.font.size = Pt(sizes.get(level, 12))
    if level == 1:
        run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    elif level == 2:
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    elif level == 3:
        run.font.color.rgb = RGBColor(0x2F, 0x84, 0x6E)


def add_paragraph_with_inline(doc, text, style="Normal"):
    """Přidá odstavec s podporou **tučného** a `kódového` textu."""
    para = doc.add_paragraph(style=style)
    # Rozdělení na segmenty: **bold**, `code`, normální text
    pattern = re.compile(r'(\*\*[^*]+\*\*|`[^`]+`)')
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = para.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = para.add_run(part[1:-1])
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            para.add_run(part)
    return para


def add_code_block(doc, lines):
    """Přidá blok kódu s monospaed fontem a šedým pozadím (přes XML)."""
    code_text = "\n".join(lines)
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after = Pt(4)
    para.paragraph_format.left_indent = Cm(0.5)

    # Šedé pozadí přes pPr shading
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)

    run = para.add_run(code_text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    return para


def add_table_from_md(doc, header_line, separator_line, data_lines):
    cols = [c.strip() for c in header_line.strip("|").split("|")]
    num_cols = len(cols)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, col in enumerate(cols):
        hdr_cells[i].text = col
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        # Šedé pozadí záhlaví
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "D6E4F0")
        tcPr.append(shd)

    for line in data_lines:
        row_cells = table.add_row().cells
        cells = [c.strip() for c in line.strip("|").split("|")]
        for i in range(min(num_cols, len(cells))):
            row_cells[i].text = cells[i]

    doc.add_paragraph()


def convert():
    doc = Document()

    # --- Styl dokumentu ---
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    with open(MD_FILE, encoding="utf-8") as f:
        lines = f.readlines()

    in_code = False
    code_lines = []
    table_header = None
    table_sep = None
    table_rows = []
    in_table = False
    skip_toc = False  # přeskočíme sekci Obsah (markdown-style interní odkazy)

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # --- Kódový blok ---
        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                in_code = False
                add_code_block(doc, code_lines)
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        # --- Tabulka ---
        if line.startswith("|"):
            if table_header is None:
                table_header = line
            elif re.match(r"^\|[\s\-:|]+\|", line):
                table_sep = line
            else:
                table_rows.append(line)
            i += 1
            # Podívej se, zda další řádek je stále tabulka
            if i < len(lines) and lines[i].startswith("|"):
                continue
            else:
                # Konec tabulky
                if table_header and table_sep:
                    add_table_from_md(doc, table_header, table_sep, table_rows)
                table_header = None
                table_sep = None
                table_rows = []
                continue

        # --- Nadpisy ---
        heading_match = re.match(r"^(#{1,4})\s+(.*)", line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Přeskočení obsahu (seznam interních odkazů)
            if text in ("Obsah",):
                skip_toc = True
                i += 1
                continue
            skip_toc = False
            heading_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
            para = doc.add_heading(text, level=level)
            set_heading_style(para, level)
            i += 1
            continue

        # --- Přeskočení obsahu (číslované položky s markdown linky) ---
        if skip_toc and re.match(r"^\s*\d+\.", line):
            i += 1
            continue
        if skip_toc and re.match(r"^\s+[-*]", line):
            i += 1
            continue
        if skip_toc and line.strip() == "":
            # prázdný řádek ukončí přeskakování
            skip_toc = False
            i += 1
            continue
        if skip_toc and line.strip() == "---":
            skip_toc = False
            i += 1
            continue

        # --- Horizontální čára ---
        if re.match(r"^-{3,}$", line.strip()):
            doc.add_paragraph("─" * 60)
            i += 1
            continue

        # --- Odrážky ---
        bullet_match = re.match(r"^(\s*)[-*]\s+(.*)", line)
        if bullet_match:
            indent = len(bullet_match.group(1)) // 2
            text = bullet_match.group(2)
            para = add_paragraph_with_inline(doc, text, style="List Bullet")
            para.paragraph_format.left_indent = Cm(0.5 + indent * 0.5)
            i += 1
            continue

        # --- Číslované odrážky ---
        numbered_match = re.match(r"^\s*\d+\.\s+(.*)", line)
        if numbered_match:
            text = numbered_match.group(1)
            add_paragraph_with_inline(doc, text, style="List Number")
            i += 1
            continue

        # --- Prázdný řádek ---
        if line.strip() == "":
            doc.add_paragraph()
            i += 1
            continue

        # --- Normální odstavec ---
        add_paragraph_with_inline(doc, line)
        i += 1

    doc.save(OUT_FILE)
    print(f"Hotovo: {OUT_FILE}")


if __name__ == "__main__":
    convert()
